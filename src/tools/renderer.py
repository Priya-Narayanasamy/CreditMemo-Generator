"""Rendering the memo.

Jinja2 owns all structure: section order, headings, tables, labels and every
conditional sentence. The narrative sections are the only text a model produced,
and they are dropped into slots the template controls.

The rule this module enforces is the one the whole system rests on: a value can
only reach the memo through `f()`, which reads the evidence ledger and raises
`UnsourcedFigure` if the field is not there. A missing ledger key fails the render
rather than quietly producing a blank cell, because a blank cell in a credit memo
reads as a fact.
"""

from __future__ import annotations

from datetime import date, datetime, timezone
from decimal import Decimal
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, StrictUndefined

from config import POLICY_VERSION, TEMPLATES_DIR
from src.state import MemoState
from src.tools.calculators import format_cents, format_ratio_as_percent
from src.tools.models import model_identifiers


class UnsourcedFigure(KeyError):
    """A template asked for a value that is not in the evidence ledger.

    This is deliberately fatal. There is no rendering path that produces a memo
    with an unsourced figure in it.
    """


PRODUCT_LABELS = {
    "variable": "Variable rate",
    "fixed_3yr": "Fixed, 3 years",
    "interest_only": "Interest only",
}

STATUS_LABELS = {
    "within_parameter": "Within parameter",
    "outside_parameter": "Outside parameter",
    "not_evaluable": "Not evaluable",
}

PURPOSE_LABELS = {
    "owner_occupied": "Owner occupied",
    "investment": "Investment",
}

STRUCTURE_LABELS = {
    "single": "Single applicant",
    "joint": "Joint applicants",
}

VERIFICATION_LABELS = {
    "full_name": "Name",
    "date_of_birth": "Date of birth",
    "residential_address": "Residential address",
}

OUTCOME_LABELS = {
    "match": "Agrees",
    "conflict": "Sources disagree",
    "not_present": "No document to compare",
}


def _format_date(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, str):
        try:
            value = date.fromisoformat(value)
        except ValueError:
            return value
    if isinstance(value, (date, datetime)):
        return value.strftime("%d %B %Y")
    return str(value)


def _format_money(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, int):
        return format_cents(value)
    return str(value)


def _format_percent(value: Any) -> str:
    if value is None:
        return "-"
    return format_ratio_as_percent(Decimal(str(value)))


def build_context(state: MemoState, approved_by: str | None = None) -> dict[str, Any]:
    """Everything the templates may see.

    Note what is not here: no raw computation helpers, and no way to reach a value
    except `f`, `has` and the pre-built row builders.
    """
    ledger = state.ledger

    def f(field_name: str):
        item = ledger.get(field_name)
        if item is None:
            raise UnsourcedFigure(
                f"{field_name!r} is not in the evidence ledger for "
                f"{state.application_number}. A memo may not state a value that has "
                f"no resolved source. Guard the slot with has() or escalate the field."
            )
        return item.value

    def has(field_name: str) -> bool:
        return field_name in ledger

    def narrative(section: str) -> str:
        return state.draft_sections.get(section, "").strip()

    def verification_rows(position: int) -> list[dict[str, str]]:
        rows = []
        for verification in state.verifications:
            prefix = f"borrower_{position}_"
            if not verification.field_name.startswith(prefix):
                continue
            suffix = verification.field_name[len(prefix):]
            if suffix not in VERIFICATION_LABELS:
                continue

            rows.append({
                "label": VERIFICATION_LABELS[suffix],
                "record": _cell(verification.record_value),
                "document": _cell(verification.document_value),
                "outcome": OUTCOME_LABELS[verification.status],
            })
        return rows

    def income_source_count(position: int) -> int:
        item = ledger.get(f"borrower_{position}_payg_income")
        if item is None:
            return 0
        return len(item.provenance.detail.get("inputs", []))

    borrower_positions = [
        position for position in range(1, 3)
        if f"borrower_{position}_record_name" in ledger
    ]

    unresolved = []
    for entry in sorted(state.unresolved.values(), key=lambda e: e.field_name):
        summary = "; ".join(
            f"{attempt.source_ref}: {attempt.outcome}" for attempt in entry.attempts
        ) or "no source was attempted"
        unresolved.append(
            type("UnresolvedView", (), {
                "field_name": entry.field_name,
                "reason": entry.reason,
                "conflicting_values": entry.conflicting_values,
                "attempt_summary": summary,
            })
        )

    source_counts = {"database": 0, "document": 0, "computed": 0, "analyst": 0}
    for item in ledger.values():
        source_counts[item.provenance.source_kind] += 1

    return {
        # accessors
        "f": f,
        "has": has,
        "narrative": narrative,
        "verification_rows": verification_rows,
        "income_source_count": income_source_count,
        # formatters
        "money": _format_money,
        "date": _format_date,
        "percent": _format_percent,
        "product_label": lambda value: PRODUCT_LABELS.get(value, value),
        "status_label": lambda value: STATUS_LABELS.get(value, value),
        # facts about the run
        "application_number": state.application_number,
        "template_id": state.template_id,
        "borrower_positions": borrower_positions,
        "borrower_count": len(borrower_positions),
        "purpose_label": PURPOSE_LABELS.get(
            ledger["loan_purpose"].value if "loan_purpose" in ledger else "", "-"),
        "structure_label": STRUCTURE_LABELS.get(
            ledger["applicant_structure"].value if "applicant_structure" in ledger else "", "-"),
        "buffer_label": _buffer_label(state),
        "policy_findings": state.policy_findings,
        "documents": state.documents,
        "unresolved": unresolved,
        "ledger_size": len(ledger),
        "database_count": source_counts["database"],
        "document_count": source_counts["document"],
        "computed_count": source_counts["computed"],
        "policy_version": POLICY_VERSION,
        "models": model_identifiers(),
        "generated_at": datetime.now(timezone.utc).strftime("%d %B %Y %H:%M UTC"),
        "approved_by": approved_by or state.approved_by or "not yet approved",
    }


def _cell(item) -> str:
    if item is None:
        return "not found"
    return str(item.value)


def _buffer_label(state: MemoState) -> str:
    item = state.ledger.get("computed_assessed_repayments")
    if item is None:
        return "-"
    buffer = item.provenance.detail.get("assessment_rate_buffer")
    return f"{Decimal(buffer) * 100:.0f}% uplift on the scheduled repayment" if buffer else "-"


def environment(templates_dir: Path | None = None) -> Environment:
    """Strict undefined, so a typo in a template is a failure and not a blank."""
    return Environment(
        loader=FileSystemLoader(str(templates_dir or TEMPLATES_DIR)),
        undefined=StrictUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
        keep_trailing_newline=True,
    )


def render_markdown(state: MemoState, approved_by: str | None = None) -> str:
    """Render the memo to Markdown. Raises rather than emitting an unsourced value."""
    if not state.template_id:
        raise ValueError("no template selected; the evidence loop has not run")

    template = environment().get_template(f"{state.template_id}.md.j2")
    return template.render(**build_context(state, approved_by))


# --- docx -------------------------------------------------------------------


def render_docx(state: MemoState, path: Path, approved_by: str | None = None) -> Path:
    """Write the memo as a .docx.

    The footer records the application number, the policy ruleset version, the
    model identifiers, the generation timestamp and the approving user.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    markdown = render_markdown(state, approved_by)
    context = build_context(state, approved_by)

    document = Document()
    _write_blocks(document, markdown)

    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = (
        f"{context['application_number']} | policy {context['policy_version']} | "
        f"extraction {context['models']['extraction']} | "
        f"drafting {context['models']['drafting']} | "
        f"review {context['models']['review']} | "
        f"generated {context['generated_at']} | "
        f"approved by {context['approved_by']}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in footer.runs:
        run.font.size = Pt(7)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


def _write_blocks(document, markdown: str) -> None:
    """Turn the rendered Markdown into docx blocks.

    Deliberately small: the templates only ever emit headings, paragraphs, bullets
    and pipe tables, so a full Markdown parser would be more surface than this
    needs.
    """
    lines = markdown.splitlines()
    index = 0
    paragraph: list[str] = []

    def flush():
        nonlocal paragraph
        if paragraph:
            document.add_paragraph(" ".join(paragraph))
            paragraph = []

    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("|"):
            flush()
            rows = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(set(cell) <= set("-: ") for cell in cells):
                    rows.append(cells)
                index += 1
            _write_table(document, rows)
            continue

        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            document.add_heading(line.lstrip("# ").strip(), level=min(level, 4))
        elif line.startswith(("- ", "* ")):
            flush()
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("---"):
            flush()
            document.add_page_break()
        elif not line:
            flush()
        else:
            paragraph.append(line)

        index += 1

    flush()


def _write_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return

    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Table Grid"

    for position, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(width):
            cells[column].text = row[column] if column < len(row) else ""
            if position == 0:
                for paragraph in cells[column].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
